import os
import sys
import zipfile
from docx import Document
import pymupdf
from io import BytesIO
from PIL import Image
import chardet
from cross_gpt import let_log
import gc
import numpy as np
import pandas as pd
import cv2
import unicodedata
from cross_gpt import (
    err_image_process_text_infoloaders,
    text_on_image_prompt_infoloaders,
    err_image_process_pdf_infoloaders,
    page_pdf_prompt_infoloaders,
    attachment_prompt_infoloaders,
    unprocessable_file_infoloaders,
    file_processing_error_infoloaders,
    image_processing_error_infoloaders,
    corrupted_zip_infoloaders,
    zip_processing_error_infoloaders,
    unsupported_format_infoloaders,
    file_open_error_infoloaders,
    zip_archive_name_infoloaders,
    model_early_loading_error_text,
    excel_cheet_text,
    excel_cheet_size_text,
    excel_cheet_strings_text,
    excel_cheet_columns_text,
    excel_cheet_data_text,
    excel_cheet_error_text,
    excel_empty_text,
    excel_error_text,
)

def get_resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'): return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

def get_external_path(relative_path):
    """Получает путь к файлу/папке рядом с .exe или скриптом"""
    if hasattr(sys, '_MEIPASS'):
        # Если запущено как exe, берем директорию, где лежит сам exe
        return os.path.join(os.path.dirname(sys.executable), relative_path)
    # Если запущен просто .py, берем текущую папку
    return os.path.join(os.path.abspath("."), relative_path)

# Глобальные переменные для ленивой загрузки моделей
_BLIP_PROCESSOR = None
_BLIP_MODEL = None
_BLIP_TOKENIZER = None
_CLIP_PROCESSOR = None
_CLIP_MODEL = None
_OCR_INSTANCE = None
_IMAGE_MODELS_LOADED = False
_IMAGE_MODELS_LOAD_FAILED = False  # Маркер ошибки загрузки моделей

def _load_image_models():
    """Ленивая загрузка моделей обработки изображений из внешней папки data/models/"""
    global _BLIP_PROCESSOR, _BLIP_MODEL, _OCR_INSTANCE
    global _IMAGE_MODELS_LOADED, _IMAGE_MODELS_LOAD_FAILED
    if _IMAGE_MODELS_LOADED: return True
    if _IMAGE_MODELS_LOAD_FAILED: raise RuntimeError(model_early_loading_error_text)
    let_log("Загрузка моделей обработки изображений из локальных директорий...")
    try:
        import torch # попытка решения проблемы dll
        import os
        from transformers import BlipProcessor, BlipForConditionalGeneration
        import easyocr
        # Определяем пути к внешним папкам
        # Ожидаемая структура: [папка с exe]/data/models/blip и easyocr
        blip_path = get_external_path(os.path.join("data", "models", "blip"))
        easyocr_path = get_external_path(os.path.join("data", "models", "easyocr"))
        # 1. Загрузка BLIP
        if os.path.exists(blip_path):
            _BLIP_PROCESSOR = BlipProcessor.from_pretrained(blip_path, use_fast=True)
            _BLIP_MODEL = BlipForConditionalGeneration.from_pretrained(blip_path)
            let_log("BLIP загружен локально")
        else:
            let_log(f"ВНИМАНИЕ: Локальная модель BLIP не найдена в {blip_path}. Попытка загрузки из сети...")
            _BLIP_PROCESSOR = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base", use_fast=True)
            _BLIP_MODEL = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
        # 2. Загрузка EasyOCR
        try:
            # Создаем папку, если её нет, чтобы Reader не ругался
            os.makedirs(easyocr_path, exist_ok=True)
            _OCR_INSTANCE = easyocr.Reader(
                ['en', 'ru'], 
                gpu=False,
                model_storage_directory=easyocr_path,
                download_enabled=False # Строго берем только из папки
            )
            let_log("EasyOCR загружен локально")
        except Exception as ocr_error:
            let_log(f"Ошибка загрузки EasyOCR из {easyocr_path}: {ocr_error}")
            _OCR_INSTANCE = None
        _IMAGE_MODELS_LOADED = True
        _IMAGE_MODELS_LOAD_FAILED = False
        return True
    except Exception as e:
        let_log(f"Критическая ошибка загрузки моделей: {e}")
        _IMAGE_MODELS_LOADED = False
        _IMAGE_MODELS_LOAD_FAILED = True
        raise

def is_binary_text(text, threshold=0.5, max_chars=1024):
    """
    Проверяет, является ли текст бинарным (содержит много непечатаемых символов).
    
    Args:
        text (str): Текст для проверки.
        threshold (float): Пороговое значение доли печатаемых символов.
        max_chars (int): Максимальное количество символов для проверки.
    
    Returns:
        bool: True, если текст считается бинарным.
    """
    if not text: return True
    check_text = text[:max_chars]
    printable_count = 0
    for char in check_text:
        category = unicodedata.category(char)
        if char == '\n' or char == '\r' or char == '\t': printable_count += 1
        elif category in ('Cc', 'Cs'): continue
        elif char == '\ufffd': continue
        else:
            printable_count += 1
    ratio = printable_count / len(check_text)
    return ratio < threshold

def process_image(file_path_or_data, input_file_handlers):
    """Обработка изображения с использованием BLIP и EasyOCR"""
    let_log('Обработка изображения')
    try: _load_image_models()
    except Exception as e: return f"Ошибка загрузки моделей: {e}"
    try:
        # Определяем тип входных данных
        if isinstance(file_path_or_data, bytes):
            let_log('  Обработка изображения из байтов (из архива)')
            image = Image.open(BytesIO(file_path_or_data)).convert('RGB')
        else:
            let_log(f'  Обработка изображения из файла: {file_path_or_data}')
            image = Image.open(file_path_or_data).convert('RGB')
    except Exception as e: return f"{err_image_process_text_infoloaders}{e}"
    # BLIP анализ (описание картинки)
    try:
        inputs = _BLIP_PROCESSOR(image, return_tensors="pt")
        output_ids = _BLIP_MODEL.generate(**inputs, max_length=50, num_beams=4)
        caption = _BLIP_PROCESSOR.decode(output_ids[0], skip_special_tokens=True)
    except Exception as e:
        let_log(f"BLIP ошибка: {e}")
        caption = "Не удалось создать описание"
    # EasyOCR обработка (извлечение текста)
    extracted_text = ""
    if _OCR_INSTANCE is not None:
        try:
            # Преобразуем PIL Image в numpy array для EasyOCR
            img_np = np.array(image)
            # Вызываем специфичный для EasyOCR метод .readtext()
            results = _OCR_INSTANCE.readtext(img_np)
            # Извлекаем только текст (он находится во втором элементе кортежа: (bbox, text, prob))
            extracted_text = "\n".join([res[1] for res in results])
        except Exception as ocr_error:
            let_log(f"EasyOCR не удался: {ocr_error}")
            extracted_text = ""
    if extracted_text: return f"{caption}\n{text_on_image_prompt_infoloaders}\n{extracted_text}"
    else: return f"{caption}"

def cleanup_image_models():
    """
    Очищает глобальные переменные с моделями и вызывает сборщик мусора
    Вызывается после обработки всех пользовательских файлов
    """
    global _BLIP_PROCESSOR, _BLIP_MODEL, _BLIP_TOKENIZER
    global _CLIP_PROCESSOR, _CLIP_MODEL, _OCR_INSTANCE
    global _IMAGE_MODELS_LOADED, _IMAGE_MODELS_LOAD_FAILED
    let_log("Очистка моделей обработки изображений...")
    # Очищаем переменные
    _BLIP_PROCESSOR = None
    _BLIP_MODEL = None
    _BLIP_TOKENIZER = None
    _CLIP_PROCESSOR = None
    _CLIP_MODEL = None
    _OCR_INSTANCE = None
    _IMAGE_MODELS_LOADED = False
    _IMAGE_MODELS_LOAD_FAILED = False  # Сбрасываем маркер ошибки при очистке
    # Принудительный вызов сборщика мусора
    gc.collect()
    let_log("Модели очищены, сборщик мусора вызван")

def process_pdf(file_path_or_data, input_file_handlers):
    let_log('пдф')
    try:
        # Определяем тип входных данных
        if isinstance(file_path_or_data, bytes):
            let_log('  Обработка PDF из байтов (из архива)')
            pdf = pymupdf.open(stream=BytesIO(file_path_or_data), filetype="pdf")
        else:
            let_log(f'  Обработка PDF из файла: {file_path_or_data}')
            pdf = pymupdf.open(file_path_or_data)
    except Exception as ex:
        let_log(ex)
        return f"{err_image_process_pdf_infoloaders}{ex}"
    full_text = []
    attachment_count = 0
    # Проход по страницам документа
    for page_num, page in enumerate(pdf, start=1):
        # Извлечение текста со страницы
        text = page.get_text()
        if text.strip():
            full_text.append(f"--- {page_pdf_prompt_infoloaders} {page_num} ---")
            full_text.append(text.strip())
        # Извлечение изображений
        for image_index, img in enumerate(page.get_images(full=True), start=1):
            xref = img[0]  # Индекс изображения
            base_image = pdf.extract_image(xref)
            image_bytes = base_image["image"]
            extension = base_image["ext"]
            attachment_count += 1
            file_name = f"{attachment_prompt_infoloaders}{attachment_count}"
            try:
                # Передаем изображение как байты
                result = input_file_handlers.get(extension, lambda x: unprocessable_file_infoloaders)(image_bytes, input_file_handlers)
                full_text.append(f"[{file_name}: {result}]")
            except Exception as ex:
                let_log(ex)
                full_text.append(f"[{file_name}: {image_processing_error_infoloaders} ({ex})]")
    pdf.close()
    return "\n".join(full_text)

def process_docx(file_path_or_data, input_file_handlers):
    let_log('док икс')
    try:
        # Определяем тип входных данных
        if isinstance(file_path_or_data, bytes):
            # Если переданы байты, используем BytesIO
            let_log('  Обработка DOCX из байтов (из архива)')
            doc = Document(BytesIO(file_path_or_data))
        else:
            # Иначе считаем, что это путь к файлу
            let_log(f'  Обработка DOCX из файла: {file_path_or_data}')
            doc = Document(file_path_or_data)
    except Exception as e: return f"{file_open_error_infoloaders}{e}"
    full_text = []
    attachment_count = 0
    # Проход по всем элементам документа
    for paragraph in doc.paragraphs:
        # Добавляем текст из параграфа
        if paragraph.text.strip(): full_text.append(paragraph.text.strip())
        # Проверяем наличие изображений в параграфе
        for run in paragraph.runs:
            # Ищем встроенные изображения
            if hasattr(run, 'element') and hasattr(run.element, 'xpath'):
                try:
                    # Ищем теги blip (встроенные изображения)
                    embed_elements = run.element.xpath('.//a:blip/@r:embed')
                    if embed_elements:
                        for embed_id in embed_elements:
                            attachment_count += 1
                            try:
                                # Получаем часть документа, соответствующую медиафайлу
                                if hasattr(doc.part, 'related_parts') and embed_id in doc.part.related_parts:
                                    media_part = doc.part.related_parts[embed_id]
                                    # Определяем расширение файла
                                    if hasattr(media_part, 'partname'): extension = os.path.splitext(media_part.partname)[-1]
                                    else: extension = '.jpg'  # По умолчанию
                                    file_name_with_ext = f"{attachment_prompt_infoloaders}{attachment_count}{extension}"
                                    # Получаем содержимое изображения
                                    if hasattr(media_part, 'blob'):
                                        image_content = media_part.blob
                                        # Определяем расширение для обработчика
                                        ext_for_handler = extension[1:] if extension.startswith('.') else extension
                                        # Проверяем, есть ли обработчик для этого расширения
                                        if ext_for_handler in input_file_handlers:
                                            # Передаём содержимое изображения в обработчик
                                            result = input_file_handlers[ext_for_handler](image_content, input_file_handlers)
                                            full_text.append(f"[{file_name_with_ext}: {result}]")
                                        else: full_text.append(f"[{file_name_with_ext}: {unsupported_format_infoloaders}]")
                                    else: full_text.append(f"[{file_name_with_ext}: Нет данных изображения]")
                                else: full_text.append(f"[Изображение {attachment_count}: Не найдено в документе]")
                            except Exception as e: full_text.append(f"[Изображение {attachment_count}: {file_processing_error_infoloaders} ({e})]")
                except Exception as e:
                    # Игнорируем ошибки при поиске изображений
                    let_log(f"  Ошибка при поиске изображений в DOCX: {e}")
                    continue
    return "\n".join(full_text)

def process_zip(file_path_or_data, input_file_handlers, is_nested=False, depth=0, max_depth=5):
    """
    Обрабатывает ZIP-архив с рекурсивной обработкой вложенных архивов
    Args:
        file_path_or_data: путь к файлу или байты архива
        input_file_handlers: словарь обработчиков файлов
        is_nested: флаг вложенного архива (для логирования)
        depth: текущая глубина рекурсии
        max_depth: максимальная глубина рекурсии (защита от бесконечной вложенности)
    """
    if depth >= max_depth:
        let_log(f"⚠ Достигнута максимальная глубина рекурсии ({max_depth}) для архивов")
        return [{
            'filename': zip_archive_name_infoloaders,
            'content': f"Достигнута максимальная глубина вложенности архивов ({max_depth})",
            'type': 'error'
        }]
    let_log('ZIP-архив' + (' (вложенный)' if is_nested else ''))
    results = []
    try:
        # Определяем тип входных данных
        if isinstance(file_path_or_data, bytes): zip_file = zipfile.ZipFile(BytesIO(file_path_or_data))
        else: zip_file = zipfile.ZipFile(file_path_or_data)
        with zip_file:
            for file_info in zip_file.infolist():
                # Пропускаем директории
                if file_info.is_dir(): continue
                # Получаем полный путь к файлу (включая подпапки)
                file_path = file_info.filename
                # Получаем расширение файла
                _, ext = os.path.splitext(file_info.filename)
                ext = ext[1:].lower()  # Убираем точку и приводим к нижнему регистру
                # Читаем содержимое файла
                try: file_content = zip_file.read(file_info)
                except Exception as read_error:
                    results.append({
                        'filename': file_path,
                        'content': f"Ошибка чтения файла из архива: {read_error}",
                        'type': 'error'
                    })
                    continue
                # Проверяем, является ли файл ZIP-архивом (рекурсивная обработка)
                if ext == 'zip':
                    let_log(f"  Обнаружен вложенный архив: {file_path}")
                    try:
                        # Рекурсивная обработка вложенного архива
                        nested_results = process_zip(
                            file_content, 
                            input_file_handlers, 
                            is_nested=True,
                            depth=depth + 1,
                            max_depth=max_depth
                        )
                        for nested_result in nested_results:
                            # Добавляем путь родительского архива к имени файла
                            nested_result['filename'] = f"{file_path}/{nested_result['filename']}"
                            results.append(nested_result)
                    except Exception as e:
                        results.append({
                            'filename': file_path,
                            'content': f"Ошибка обработки вложенного архива: {e}",
                            'type': 'error'
                        })
                    continue
                # Обрабатываем обычные файлы соответствующим обработчиком
                if ext in input_file_handlers:
                    try:
                        result = input_file_handlers[ext](file_content, input_file_handlers)
                        results.append({
                            'filename': file_path,
                            'content': result,
                            'type': 'file'
                        })
                    except Exception as e:
                        results.append({
                            'filename': file_path,
                            'content': f"{file_processing_error_infoloaders}{e}",
                            'type': 'error'
                        })
                else:
                    # Если формат не поддерживается, просто добавляем информацию о файле
                    results.append({
                        'filename': file_path,
                        'content': unsupported_format_infoloaders,
                        'type': 'unsupported'
                    })
    except zipfile.BadZipFile:
        error_msg = corrupted_zip_infoloaders
        if is_nested: error_msg = f"Вложенный архив поврежден: {error_msg}"
        return [{'filename': zip_archive_name_infoloaders, 'content': error_msg, 'type': 'error'}]
    except Exception as e:
        error_msg = f'{zip_processing_error_infoloaders}: {e}'
        if is_nested: error_msg = f"Ошибка обработки вложенного архива: {error_msg}"
        return [{'filename': zip_archive_name_infoloaders, 'content': error_msg, 'type': 'error'}]
    return results

def process_text(file_path_or_data, input_file_handlers):
    let_log('text')
    try:
        if isinstance(file_path_or_data, bytes):
            let_log('  Обработка текста из байтов (из архива)')
            raw_data = file_path_or_data
        else:
            let_log(f'  Обработка текста из файла: {file_path_or_data}')
            try:
                with open(file_path_or_data, 'rb') as file: raw_data = file.read()
            except (FileNotFoundError, IOError):
                # Если файл не найден, считаем что это сам текст
                return file_path_or_data
        # Определяем кодировку
        result = chardet.detect(raw_data)
        encoding = result['encoding'] or 'utf-8'  # Защита от None
        # Декодируем с обработкой ошибок
        return raw_data.decode(encoding, errors='replace')
    except Exception as e:
        let_log(f"  Ошибка обработки текста: {e}")
        # Пытаемся вернуть как есть
        if isinstance(file_path_or_data, bytes): return file_path_or_data.decode('utf-8', errors='ignore')
        else: return file_path_or_data

def process_excel(file_path_or_data, input_file_handlers):
    """
    Обрабатывает Excel файлы (xlsx, xls) - извлекает данные из всех листов
    Поддерживает как путь к файлу, так и байты
    """
    let_log('Excel файл')
    try:
        # Определяем тип входных данных
        if isinstance(file_path_or_data, bytes):
            let_log('  Обработка Excel из байтов (из архива)')
            # Используем BytesIO для чтения из памяти
            import io
            excel_file = pd.ExcelFile(io.BytesIO(file_path_or_data))
        else:
            let_log(f'  Обработка Excel из файла: {file_path_or_data}')
            excel_file = pd.ExcelFile(file_path_or_data)
        results = []
        # Проходим по всем листам
        for sheet_name in excel_file.sheet_names:
            try:
                # Читаем лист в DataFrame
                df = excel_file.parse(sheet_name)
                # Преобразуем DataFrame в текстовое представление
                df_text = df.fillna('').astype(str).to_string(index=False)
                results.append(f"--- {excel_cheet_text}{sheet_name} ---")
                results.append(f"{excel_cheet_size_text}{df.shape[0]} {excel_cheet_strings_text} {df.shape[1]} {excel_cheet_columns_text}")
                results.append(excel_cheet_data_text)
                results.append(df_text)
                results.append("")  # Пустая строка для разделения
            except Exception as e:
                results.append(f"{excel_cheet_error_text} '{sheet_name}' {e}")
                continue
        excel_file.close()
        if results: return "\n".join(results)
        else: return excel_empty_text
    except Exception as e:
        error_msg = f"{excel_error_text} {e}"
        let_log(error_msg)
        return error_msg

def process_unknown(file_path_or_data, input_file_handlers):
    """
    Обработчик для файлов с неизвестными расширениями.
    Пытается обработать файл как текстовый.
    """
    let_log('Попытка обработки неизвестного файла как текстового')
    text = process_text(file_path_or_data, input_file_handlers)
    if is_binary_text(text): raise ValueError("Файл не является текстовым или имеет неизвестную кодировку.")
    else:
        let_log('Файл обработан как текстовый')
        return text